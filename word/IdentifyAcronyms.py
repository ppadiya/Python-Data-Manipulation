import sys
sys.path.insert(0, './python-docx-master')
from docx import Document

import docx
import re

def find_acronyms(doc):
    acronyms = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Use a regular expression to find acronyms (capital letters)
        matches = re.findall(r'\b[A-Z]+\b', text)
        acronyms.update(matches)
    return acronyms

def add_acronym_list(doc, acronyms):
    doc.add_paragraph("List of Acronyms:")
    for acronym in acronyms:
        doc.add_paragraph(f"{acronym}", style='List Bullet')

def main(input_file, output_file):
    try:
        doc = docx.Document(input_file)
        acronyms = find_acronyms(doc)
        if acronyms:
            add_acronym_list(doc, acronyms)
            doc.save(output_file)
            print("Acronyms added to the document.")
        else:
            print("No acronyms found in the document.")
    except Exception as e:
        print(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    input_file = "input.docx"  # Replace with your input file name
    output_file = "output.docx"  # Replace with your output file name
    main(input_file, output_file)
